import docx

doc = docx.Document("CUET Thesis Template in Word v1(3).docx")

print("=== Scanning Paragraphs 40 to 70 ===")
for i in range(40, 71):
    txt = doc.paragraphs[i].text.strip()
    if txt:
        print(f"P{i:03d} [Style: {doc.paragraphs[i].style.name}]: {txt}")

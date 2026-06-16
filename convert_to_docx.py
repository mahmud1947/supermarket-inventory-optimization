import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner margins (padding) for a table cell in dxa (1 dxa = 1/20 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, color="CCCCCC", sz="4", val="single"):
    """Set thin borders for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right']:
        node = OxmlElement(border_name)
        node.set(qn('w:val'), val)
        node.set(qn('w:sz'), sz)
        node.set(qn('w:space'), '0')
        node.set(qn('w:color'), color)
        tcBorders.append(node)
    tcPr.append(tcBorders)

def parse_runs(paragraph, text, default_font="Georgia"):
    """
    Parses bold (**), italic (*), and inline math ($) patterns.
    Applies appropriate run properties (bold, italic, sub/sup) in Word.
    """
    # Pattern to tokenize: Math ($...$), Bold (**...**), Italic (*...*)
    tokens = re.split(r'(\$\$[^\$]+\$\$|\$[^\$]+\$|\*\*[^*]+\*\*|\*[^*]+\*)', text)
    
    for token in tokens:
        if not token:
            continue
            
        # Math Block ($$ ... $$) or Inline Math ($ ... $)
        if token.startswith('$'):
            is_block = token.startswith('$$')
            math_content = token.strip('$ ')
            
            # Clean up latex keywords
            math_content = re.sub(r'\\(mu|sigma|lambda|beta|alpha|pi|theta|delta|Phi|Delta)', lambda m: {
                'mu': 'μ', 'sigma': 'σ', 'lambda': 'λ', 'beta': 'β', 'alpha': 'α',
                'pi': 'π', 'theta': 'θ', 'delta': 'δ', 'Phi': 'Φ', 'Delta': 'Δ'
            }[m.group(1)], math_content)
            
            # Common symbols
            math_content = math_content.replace(r'\cdot', '·')
            math_content = math_content.replace(r'\times', '×')
            math_content = math_content.replace(r'\le', '≤')
            math_content = math_content.replace(r'\ge', '≥')
            math_content = math_content.replace(r'\approx', '≈')
            math_content = math_content.replace(r'\infty', '∞')
            math_content = math_content.replace(r'\sum', 'Σ')
            math_content = math_content.replace(r'\text', '')
            math_content = math_content.replace(r'\left(', '(').replace(r'\right)', ')')
            math_content = math_content.replace(r'\left[', '[').replace(r'\right]', ']')
            math_content = re.sub(r'\\(quad|qquad|;|!|,)', ' ', math_content)
            math_content = math_content.replace('\\', '')
            
            # Parse subscripts (_) and superscripts (^)
            sub_sup_tokens = re.split(r'(_\{[^}]+\}|_[a-zA-Z0-9*+-]+|\^\{[^}]+\}|\^[a-zA-Z0-9*+-]+)', math_content)
            
            for s_token in sub_sup_tokens:
                if not s_token:
                    continue
                run = paragraph.add_run()
                run.font.name = default_font
                run.font.size = Pt(11.5)
                
                if s_token.startswith('_'):
                    clean_val = s_token[1:].strip('{}')
                    run.text = clean_val
                    run.font.subscript = True
                    run.font.italic = True
                elif s_token.startswith('^'):
                    clean_val = s_token[1:].strip('{}')
                    run.text = clean_val
                    run.font.superscript = True
                    run.font.italic = True
                else:
                    run.text = s_token
                    # Italicize math variables, but keep numbers and brackets normal
                    if re.match(r'^[a-zA-Z0-9*+-]+$', s_token) and not s_token.isdigit():
                        run.font.italic = True
                    
        # Bold (** ... **)
        elif token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.font.name = default_font
            run.font.size = Pt(11.5)
            run.font.bold = True
            
        # Italic (* ... *)
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = default_font
            run.font.size = Pt(11.5)
            run.font.italic = True
            
        # Normal Text
        else:
            run = paragraph.add_run(token)
            run.font.name = default_font
            run.font.size = Pt(11.5)

def build_docx():
    print("Reading thesis_draft.md to compile into Microsoft Word format...")
    if not os.path.exists("thesis_draft.md"):
        print("Error: thesis_draft.md not found. Please compile it first.")
        return

    with open("thesis_draft.md", "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    
    # 1. Page Setup (CUET Academic Rules: A4, Left=1.38", Right=0.98", Top/Bottom=0.98")
    for section in doc.sections:
        section.page_width = Inches(8.27)  # A4 width
        section.page_height = Inches(11.69) # A4 height
        section.top_margin = Inches(0.98)
        section.bottom_margin = Inches(0.98)
        section.left_margin = Inches(1.38)
        section.right_margin = Inches(0.98)

    # Separate cover page HTML from the rest of markdown content
    parts = content.split("</div>", 1)
    if len(parts) < 2:
        print("Error: Could not split title page from the rest of the document.")
        return
        
    cover_html = parts[0]
    rest_markdown = parts[1]

    # Render Title Page (Page 1)
    print("Rendering Title Page...")
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(24)
    run = title_p.add_run("OPTIMIZATION OF SAFETY STOCK AND REORDERING POINT FOR A MULTI ITEM SUPPLY CHAIN USING STOCHASTIC MODELS AND MONTE CARLO SIMULATION")
    run.font.name = "Georgia"
    run.font.size = Pt(16)
    run.font.bold = True
    
    # "By"
    by_p = doc.add_paragraph()
    by_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    by_p.paragraph_format.space_before = Pt(24)
    by_p.paragraph_format.space_after = Pt(12)
    run = by_p.add_run("By")
    run.font.name = "Georgia"
    run.font.size = Pt(12)
    
    # Author Details
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.paragraph_format.space_after = Pt(36)
    run1 = author_p.add_run("Md Mahmudur Rahman\n")
    run1.font.bold = True
    run1.font.size = Pt(13)
    run2 = author_p.add_run("Student ID: 2009007")
    run2.font.bold = True
    run2.font.size = Pt(12)
    for r in [run1, run2]:
        r.font.name = "Georgia"
        
    # Logo
    logo_path = "chittagong-university-of-engineering-and-technolog-seeklogo.png"
    if os.path.exists(logo_path):
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.space_after = Pt(36)
        logo_p.add_run().add_picture(logo_path, width=Inches(1.5))
        
    # Degree info
    degree_p = doc.add_paragraph()
    degree_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    degree_p.paragraph_format.space_after = Pt(36)
    run1 = degree_p.add_run("A thesis submitted in partial fulfilment of the requirements for the degree of\n")
    run1.font.size = Pt(11)
    run2 = degree_p.add_run("Bachelor of Science in Mechatronics and Industrial Engineering")
    run2.font.bold = True
    run2.font.size = Pt(12)
    for r in [run1, run2]:
        r.font.name = "Georgia"
        
    # Department & University
    univ_p = doc.add_paragraph()
    univ_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    univ_p.paragraph_format.space_before = Pt(24)
    univ_p.paragraph_format.space_after = Pt(12)
    run1 = univ_p.add_run("Department of Mechatronics & Industrial Engineering\n")
    run1.font.bold = True
    run1.font.size = Pt(12)
    run2 = univ_p.add_run("Chittagong University of Engineering and Technology (CUET)\n")
    run2.font.bold = True
    run2.font.size = Pt(13)
    run3 = univ_p.add_run("Chittagong-4349, Bangladesh\n\n")
    run3.font.size = Pt(11)
    run4 = univ_p.add_run("JUNE 2026")
    run4.font.bold = True
    run4.font.size = Pt(11)
    for r in [run1, run2, run3, run4]:
        r.font.name = "Georgia"

    # Add page break after Cover Page
    doc.add_page_break()

    # Now parse the rest of the document block by block
    blocks = rest_markdown.split("\n\n")
    
    in_table = False
    table_rows = []
    in_code = False
    code_lines = []
    
    is_first_para = True
    
    page_break_headings = [
        "DECLARATION",
        "COPYRIGHT NOTICE",
        "DEDICATION",
        "LIST OF PUBLICATIONS",
        "APPROVAL/DECLARATION BY THE SUPERVISOR(S)",
        "ACKNOWLEDGEMENT",
        "ABSTRACT",
        "TABLE OF CONTENTS",
        "LIST OF FIGURES",
        "LIST OF TABLES",
        "NOMENCLATURE",
        "REFERENCES",
        "APPENDICES"
    ]

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        # Replace HTML break tags with newlines and clean up remaining HTML tags in this block
        block = re.sub(r'<br\s*/?>', '\n', block)
        # Note: If it's a structural div like page break, skip it or handle it
        if '<div' in block and 'page-break-before' in block:
            # We already handle page breaks via headings, so we can ignore this block or skip it
            continue
        
        # Clean other html tags (except images which we parse below)
        if not (block.startswith("<img") or "img src=" in block):
            block = re.sub(r'<[^>]+>', '', block).strip()

        if not block:
            continue

        # -------------------------------------------------------------
        # Code Block Start/End
        # -------------------------------------------------------------
        if block.startswith("```"):
            if not in_code:
                in_code = True
                code_lines.append(block.replace("```python", "").replace("```", ""))
            else:
                in_code = False
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.left_indent = Inches(0.4)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(8.5)
                code_lines = []
            continue
        
        if in_code:
            code_lines.append(block)
            continue

        # -------------------------------------------------------------
        # Table Start/Row Collection/Table End
        # -------------------------------------------------------------
        if block.startswith("|"):
            in_table = True
            table_rows.append(block)
            continue
        else:
            if in_table:
                write_table_to_docx(doc, table_rows)
                table_rows = []
                in_table = False

        # -------------------------------------------------------------
        # Headings
        # -------------------------------------------------------------
        if block.startswith("# "):
            heading_text = block[2:].strip().upper()
            
            # Heading 1 always starts on a new page!
            doc.add_page_break()
            
            h = doc.add_heading(level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = h.add_run(heading_text)
            run.font.name = "Georgia"
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(12)
            h.paragraph_format.keep_with_next = True
            is_first_para = True
            continue
            
        elif block.startswith("## "):
            heading_text = block[3:].strip()
            
            # Clean formatting characters from heading text
            clean_heading_text = heading_text.strip("*_#")
            
            # Page break for front matter pages
            if clean_heading_text.upper() in page_break_headings:
                doc.add_page_break()
                
            h = doc.add_heading(level=2)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = h.add_run(clean_heading_text)
            run.font.name = "Georgia"
            run.font.size = Pt(12.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(8)
            h.paragraph_format.keep_with_next = True
            is_first_para = True
            continue
            
        elif block.startswith("### "):
            heading_text = block[4:].strip()
            h = doc.add_heading(level=3)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = h.add_run(heading_text.strip("*_#"))
            run.font.name = "Georgia"
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.keep_with_next = True
            is_first_para = True
            continue

        # -------------------------------------------------------------
        # Images (Fitted plots or flowcharts)
        # -------------------------------------------------------------
        image_match = re.search(r'\[(.*?)\.(png|jpg|jpeg)\]\(file:///.*?\)', block)
        if image_match:
            img_name = f"{image_match.group(1)}.{image_match.group(2)}"
            img_path = f"visualizations/{img_name}"
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(img_path, width=Inches(4.5))
                caption_block = block.replace(image_match.group(0), "").strip()
                if caption_block:
                    cp = doc.add_paragraph()
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    c_run = cp.add_run(caption_block)
                    c_run.font.name = "Georgia"
                    c_run.font.size = Pt(9.5)
                    c_run.font.italic = True
            continue

        # -------------------------------------------------------------
        # Bullet/List Items
        # -------------------------------------------------------------
        if block.startswith("* ") or block.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            clean_text = block[2:].strip()
            parse_runs(p, clean_text)
            continue
        
        # Numbered Lists (like 1. 2. 3.)
        num_match = re.match(r'^(\d+)\.\s+(.*)', block, re.DOTALL)
        if num_match:
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(4)
            clean_text = num_match.group(2).strip()
            parse_runs(p, clean_text)
            continue

        # -------------------------------------------------------------
        # Normal Paragraph
        # -------------------------------------------------------------
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_after = Pt(8)
        
        if not is_first_para:
            p.paragraph_format.first_line_indent = Inches(0.5)
        
        is_first_para = False
        parse_runs(p, block)
        
    # If the page ended with a table, write it out
    if in_table:
        write_table_to_docx(doc, table_rows)
        table_rows = []
        in_table = False

    # Save the output file
    output_filename = "thesis_draft.docx"
    doc.save(output_filename)
    print(f"Microsoft Word document successfully created: {output_filename}")


def write_table_to_docx(doc, table_rows):
    """Parses markdown table text and adds a styled table to docx."""
    # Separate rows into cells
    parsed_grid = []
    for r in table_rows:
        cells = [c.strip() for c in r.split("|")[1:-1]]
        # Skip separator rows like |---|---|
        if cells and all(c == "" or re.match(r'^[-:]+$', c) for c in cells):
            continue
        if cells:
            parsed_grid.append(cells)
            
    if not parsed_grid:
        return
        
    num_rows = len(parsed_grid)
    num_cols = len(parsed_grid[0])
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    for row_idx, row_cells in enumerate(parsed_grid):
        row = table.rows[row_idx]
        for col_idx, cell_value in enumerate(row_cells):
            cell = row.cells[col_idx]
            cell.text = "" # Clear default
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Format text
            # Headers have bold text
            if row_idx == 0:
                parse_runs(p, f"**{cell_value}**")
            else:
                parse_runs(p, cell_value)
                
            # Formatting
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            set_cell_borders(cell, color="444444" if row_idx == 0 else "CCCCCC", sz="4")
            
            # Font size inside table
            for run in p.runs:
                run.font.size = Pt(8.5)


if __name__ == "__main__":
    build_docx()
